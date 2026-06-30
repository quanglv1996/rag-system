"""Document loader for the RAG pipeline.

Supports loading text content from PDF, DOCX, TXT, Markdown, and HTML
files. Each loader is a strategy implementation; new formats can be added
by registering a new loader function without modifying existing code.

All loaders return plain text — the pipeline handles chunking downstream.
"""

import asyncio
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import aiofiles

from app.common.enums import DocumentType
from app.core.exception import RAGException, ValidationException
from app.core.logger import get_logger

logger = get_logger(__name__)


@dataclass
class LoadedDocument:
    """Result of loading a document.

    Attributes:
        content: Full extracted text content.
        doc_type: Detected document type.
        source: Original file path or URL.
        metadata: Extracted metadata (title, author, page count, etc.).
    """

    content: str
    doc_type: DocumentType
    source: str
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentLoader:
    """Strategy-based document loader supporting multiple formats.

    Dispatches to the correct loader based on file extension.
    Add new formats by registering them via register_loader().
    """

    # File extension → loading coroutine mapping
    _loaders: dict[str, str] = {
        ".pdf": "_load_pdf",
        ".docx": "_load_docx",
        ".doc": "_load_docx",
        ".txt": "_load_text",
        ".md": "_load_markdown",
        ".markdown": "_load_markdown",
        ".html": "_load_html",
        ".htm": "_load_html",
    }

    async def load(self, path: str | Path) -> LoadedDocument:
        """Load a document from a file path and extract its text.

        Args:
            path: Path to the document file.

        Returns:
            LoadedDocument: Extracted content and metadata.

        Raises:
            ValidationException: If the file format is unsupported.
            RAGException: If loading fails.
        """
        file_path = Path(path)

        if not file_path.exists():
            raise RAGException(
                f"Document file not found: {path}",
                stage="load",
            )

        extension = file_path.suffix.lower()

        if extension not in self._loaders:
            raise ValidationException(
                f"Unsupported document format: '{extension}'. "
                f"Supported: {', '.join(self._loaders.keys())}",
                field="file",
            )

        loader_method = getattr(self, self._loaders[extension])

        logger.info("Loading document", path=str(path), extension=extension)

        try:
            return await loader_method(file_path)
        except (RAGException, ValidationException):
            raise
        except Exception as exc:
            raise RAGException(
                f"Failed to load document '{path}': {exc}",
                stage="load",
            ) from exc

    async def load_bytes(
        self,
        content: bytes,
        filename: str,
        metadata: dict[str, Any] | None = None,
    ) -> LoadedDocument:
        """Load a document from raw bytes (e.g., uploaded file).

        Args:
            content: Raw file bytes.
            filename: Original filename with extension.
            metadata: Optional pre-populated metadata.

        Returns:
            LoadedDocument: Extracted content and metadata.

        Raises:
            ValidationException: If the file format is unsupported.
            RAGException: If loading fails.
        """
        import tempfile

        suffix = Path(filename).suffix.lower()
        if suffix not in self._loaders:
            raise ValidationException(
                f"Unsupported document format: '{suffix}'",
                field="filename",
            )

        # Write to temp file, load, then clean up
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(content)
            tmp_path = Path(tmp.name)

        try:
            doc = await self.load(tmp_path)
            doc.source = filename
            if metadata:
                doc.metadata.update(metadata)
            return doc
        finally:
            tmp_path.unlink(missing_ok=True)

    async def _load_pdf(self, path: Path) -> LoadedDocument:
        """Load and extract text from a PDF file.

        Args:
            path: Path to the PDF file.

        Returns:
            LoadedDocument: Extracted text and page metadata.
        """
        def _extract_pdf() -> tuple[str, dict[str, Any]]:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            pages: list[str] = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)

            meta: dict[str, Any] = {
                "page_count": len(reader.pages),
                "title": reader.metadata.title if reader.metadata else None,
                "author": reader.metadata.author if reader.metadata else None,
            }

            return "\n\n".join(pages), meta

        content, metadata = await asyncio.get_event_loop().run_in_executor(
            None, _extract_pdf
        )

        return LoadedDocument(
            content=content,
            doc_type=DocumentType.PDF,
            source=str(path),
            metadata=metadata,
        )

    async def _load_docx(self, path: Path) -> LoadedDocument:
        """Load and extract text from a DOCX file.

        Args:
            path: Path to the DOCX file.

        Returns:
            LoadedDocument: Extracted text.
        """
        def _extract_docx() -> tuple[str, dict[str, Any]]:
            from docx import Document

            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            core_props = doc.core_properties

            meta: dict[str, Any] = {
                "title": core_props.title,
                "author": core_props.author,
                "paragraph_count": len(paragraphs),
            }

            return "\n\n".join(paragraphs), meta

        content, metadata = await asyncio.get_event_loop().run_in_executor(
            None, _extract_docx
        )

        return LoadedDocument(
            content=content,
            doc_type=DocumentType.DOCX,
            source=str(path),
            metadata=metadata,
        )

    async def _load_text(self, path: Path) -> LoadedDocument:
        """Load a plain text file.

        Args:
            path: Path to the text file.

        Returns:
            LoadedDocument: File content.
        """
        async with aiofiles.open(path, encoding="utf-8", errors="replace") as f:
            content = await f.read()

        return LoadedDocument(
            content=content,
            doc_type=DocumentType.TXT,
            source=str(path),
            metadata={"size_bytes": path.stat().st_size},
        )

    async def _load_markdown(self, path: Path) -> LoadedDocument:
        """Load a Markdown file and convert to plain text.

        Args:
            path: Path to the Markdown file.

        Returns:
            LoadedDocument: Plain text extracted from Markdown.
        """
        async with aiofiles.open(path, encoding="utf-8", errors="replace") as f:
            raw = await f.read()

        def _md_to_text(md_content: str) -> str:
            import re
            # Remove Markdown syntax for cleaner chunking
            text = re.sub(r"#+ ", "", md_content)           # Headers
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)  # Bold
            text = re.sub(r"\*([^*]+)\*", r"\1", text)       # Italic
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # Links
            text = re.sub(r"`[^`]+`", lambda m: m.group()[1:-1], text)  # Code
            return text

        content = await asyncio.get_event_loop().run_in_executor(
            None, _md_to_text, raw
        )

        return LoadedDocument(
            content=content,
            doc_type=DocumentType.MARKDOWN,
            source=str(path),
            metadata={},
        )

    async def _load_html(self, path: Path) -> LoadedDocument:
        """Load an HTML file and extract visible text.

        Args:
            path: Path to the HTML file.

        Returns:
            LoadedDocument: Extracted visible text.
        """
        async with aiofiles.open(path, encoding="utf-8", errors="replace") as f:
            raw = await f.read()

        def _html_to_text(html: str) -> tuple[str, str]:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(html, "lxml")
            # Remove script and style tags
            for tag in soup(["script", "style", "meta", "head"]):
                tag.decompose()

            title = soup.title.string if soup.title else ""
            text = soup.get_text(separator="\n", strip=True)
            return text, title or ""

        content, title = await asyncio.get_event_loop().run_in_executor(
            None, _html_to_text, raw
        )

        return LoadedDocument(
            content=content,
            doc_type=DocumentType.HTML,
            source=str(path),
            metadata={"title": title},
        )
